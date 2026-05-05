from __future__ import annotations

import argparse
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def clean_lines(text: str) -> list[str]:
    lines = text.splitlines()
    out: list[str] = []
    skip_style = False

    for line in lines:
        stripped = line.strip()

        if stripped.lower().startswith("<style"):
            skip_style = True
            continue
        if skip_style:
            if stripped.lower().endswith("</style>"):
                skip_style = False
            continue

        if stripped.lower().startswith("<div") or stripped.lower().startswith("</div"):
            continue

        out.append(line)

    return out


def is_table_line(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def is_separator_row(line: str) -> bool:
    s = line.strip().strip("|").replace(" ", "")
    return bool(s) and all(ch in "-:|" for ch in s)


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return [re.sub(r"\*\*(.*?)\*\*", r"\1", c) for c in cells]


def add_heading(doc: Document, line: str) -> None:
    level = len(line) - len(line.lstrip("#"))
    text = line[level:].strip()
    p = doc.add_heading(text, level=min(max(level, 1), 4))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_paragraph(doc: Document, text: str) -> None:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text).strip()
    if not text:
        doc.add_paragraph("")
        return
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_table(doc: Document, table_lines: list[str]) -> None:
    if not table_lines:
        return

    rows: list[list[str]] = []
    for line in table_lines:
        if is_separator_row(line):
            continue
        rows.append(parse_table_row(line))

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            value = row[c_idx] if c_idx < len(row) else ""
            cell = tbl.cell(r_idx, c_idx)
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = clean_lines(text)

    doc = Document()

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if is_table_line(line):
            start = i
            while i < n and is_table_line(lines[i]):
                i += 1
            add_table(doc, lines[start:i])
            continue

        if stripped.startswith("#"):
            add_heading(doc, stripped)
            i += 1
            continue

        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:].strip())
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    doc.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert simple markdown to docx")
    parser.add_argument("input", type=Path, help="Input markdown path")
    parser.add_argument("-o", "--output", type=Path, help="Output docx path")
    args = parser.parse_args()

    input_path = args.input
    output_path = args.output or input_path.with_suffix(".docx")

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    convert_md_to_docx(input_path, output_path)
    print(f"Created: {output_path}")


if __name__ == "__main__":
    main()
